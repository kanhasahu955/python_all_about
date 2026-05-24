"""Convert resume markdown to downloadable formats."""
import io
import re

from docx import Document
from docx.shared import Pt


def _add_formatted_paragraph(doc: Document, text: str, *, bullet: bool = False) -> None:
    plain = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    plain = re.sub(r"`(.+?)`", r"\1", plain).strip()
    if not plain:
        return
    if bullet:
        doc.add_paragraph(plain, style="List Bullet")
    else:
        p = doc.add_paragraph()
        run = p.add_run(plain)
        run.font.size = Pt(11)


def markdown_to_docx_bytes(markdown: str) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ", "• ")):
            _add_formatted_paragraph(doc, stripped[2:], bullet=True)
        else:
            _add_formatted_paragraph(doc, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
