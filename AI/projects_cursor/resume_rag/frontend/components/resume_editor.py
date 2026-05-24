"""Resume editor UI helpers and document export."""
import io
import re

from docx import Document
from docx.shared import Pt


def markdown_to_docx_bytes(markdown: str) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ", "• ")):
            plain = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            doc.add_paragraph(plain, style="List Bullet")
        else:
            plain = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            p = doc.add_paragraph()
            run = p.add_run(plain)
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_download_bar(content: str, file_stem: str = "optimized_resume") -> None:
    import streamlit as st

    if not content.strip():
        st.caption("Generate or edit a resume to enable downloads.")
        return

    safe = file_stem.replace(" ", "_").removesuffix(".docx").removesuffix(".md")
    col_md, col_docx, col_txt = st.columns(3)

    col_md.download_button(
        "⬇ Markdown (.md)",
        data=content,
        file_name=f"{safe}.md",
        mime="text/markdown",
        use_container_width=True,
    )
    col_docx.download_button(
        "⬇ Word (.docx)",
        data=markdown_to_docx_bytes(content),
        file_name=f"{safe}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
    col_txt.download_button(
        "⬇ Plain text (.txt)",
        data=content,
        file_name=f"{safe}.txt",
        mime="text/plain",
        use_container_width=True,
    )


def render_editor_preview(content: str, *, editor_key: str = "rb_editor") -> str:
    """Split editor + live preview; returns edited content."""
    import streamlit as st

    tab_edit, tab_preview = st.tabs(["✏️ Editor", "👁 Preview"])

    with tab_edit:
        st.caption("Edit markdown directly — changes sync to preview and downloads.")
        edited = st.text_area(
            "Resume content",
            value=content,
            height=520,
            key=editor_key,
            label_visibility="collapsed",
        )

    with tab_preview:
        st.caption("Formatted preview (ATS-style sections)")
        st.markdown(edited or content)

    return edited if edited is not None else content
